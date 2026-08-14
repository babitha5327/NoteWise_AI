"""
NOTEWISE AI - Final Production Server
======================================
Single source-of-truth Flask server. Serves the frontend, handles OCR
(project OCR service if present, else a real Tesseract pipeline with
image pre-processing), PDF/DOCX extraction, local note-processing
(summary / concepts / revision / quiz), translation, and TTS with a
browser-speech fallback.

Run:
    python3 final_server.py
Serves on http://0.0.0.0:7860
"""

import os
import io
import re
import sys
import uuid
import socket
import logging
import importlib
import traceback
from collections import Counter

from flask import Flask, request, jsonify, render_template, send_from_directory, url_for
from werkzeug.utils import secure_filename

# ---------------------------------------------------------------------------
# PATHS
# ---------------------------------------------------------------------------
BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(BACKEND_DIR)
FRONTEND_DIR = os.path.join(PROJECT_ROOT, "frontend")
TEMPLATE_DIR = os.path.join(FRONTEND_DIR, "templates")
STATIC_DIR = os.path.join(FRONTEND_DIR, "static")
DATA_DIR = os.path.join(PROJECT_ROOT, "data")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
PROCESSED_DIR = os.path.join(DATA_DIR, "processed")
AUDIO_DIR = os.path.join(PROCESSED_DIR, "audio")

for d in (UPLOAD_DIR, PROCESSED_DIR, AUDIO_DIR):
    os.makedirs(d, exist_ok=True)

logging.basicConfig(level=logging.INFO, format="[NoteWiseAI] %(message)s")
log = logging.getLogger("notewise")

# ---------------------------------------------------------------------------
# FLASK APP
# ---------------------------------------------------------------------------
app = Flask(
    __name__,
    template_folder=TEMPLATE_DIR,
    static_folder=STATIC_DIR,
    static_url_path="/static",
)
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024  # 25 MB upload limit

ALLOWED_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif", "pdf", "docx"}
IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif"}

LANGUAGES = ["English", "Tamil", "Hindi", "Telugu", "Kannada", "Malayalam"]

# ---------------------------------------------------------------------------
# OPTIONAL: detect an existing project OCR / note-pipeline service
# ---------------------------------------------------------------------------
def _try_import(path):
    try:
        module_path, func_name = path.rsplit(".", 1)
        mod = importlib.import_module(module_path)
        fn = getattr(mod, func_name, None)
        if callable(fn):
            return fn
    except Exception:
        return None
    return None

_OCR_CANDIDATES = [
    "services.ocr.process_image",
    "services.ocr_service.process_image",
    "backend.services.ocr.process_image",
    "backend.services.ocr_service.process_image",
]
_PIPELINE_CANDIDATES = [
    "services.note_pipeline.process_note",
    "services.ai_pipeline.process_note",
]

PROJECT_OCR_FN = None
for _p in _OCR_CANDIDATES:
    PROJECT_OCR_FN = _try_import(_p)
    if PROJECT_OCR_FN:
        log.info(f"Using project OCR service: {_p}")
        break

PROJECT_PIPELINE_FN = None
for _p in _PIPELINE_CANDIDATES:
    PROJECT_PIPELINE_FN = _try_import(_p)
    if PROJECT_PIPELINE_FN:
        log.info(f"Using project note pipeline: {_p}")
        break

OCR_AVAILABLE = True  # tesseract fallback is always available in this build
NOTE_PIPELINE_AVAILABLE = True  # local fallback always available

# ---------------------------------------------------------------------------
# THIRD PARTY LIBS (all optional / degrade gracefully)
# ---------------------------------------------------------------------------
try:
    import pytesseract
    from PIL import Image, ImageOps, ImageFilter, ImageEnhance
    TESSERACT_OK = True
except Exception as e:
    TESSERACT_OK = False
    log.warning(f"pytesseract/PIL unavailable: {e}")

try:
    from pypdf import PdfReader
    PYPDF_OK = True
except Exception as e:
    PYPDF_OK = False
    log.warning(f"pypdf unavailable: {e}")

try:
    import docx as docx_lib
    DOCX_OK = True
except Exception as e:
    DOCX_OK = False
    log.warning(f"python-docx unavailable: {e}")

try:
    from gtts import gTTS
    GTTS_OK = True
except Exception:
    GTTS_OK = False
    log.warning("gTTS unavailable (no network / not installed) - browser speech fallback will be used")

try:
    from deep_translator import GoogleTranslator
    TRANSLATOR_OK = True
except Exception:
    TRANSLATOR_OK = False
    log.warning("deep_translator unavailable - translation will echo original text")


# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def unique_filename(original):
    safe = secure_filename(original) or "upload"
    ext = safe.rsplit(".", 1)[-1].lower() if "." in safe else "bin"
    return f"{uuid.uuid4().hex}.{ext}"


BRAND_NOISE_PATTERNS = [
    r"note\s*wise\s*ai",
    r"handwritten\s+notes\s+to\s+intelligent\s+learning",
]


def clean_ocr_text(text):
    """Strip accidental app branding, normalise whitespace, keep line breaks."""
    if not text:
        return ""
    cleaned = text
    for pat in BRAND_NOISE_PATTERNS:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE)
    # normalise excessive blank lines but keep paragraph breaks
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = "\n".join(line.strip() for line in cleaned.split("\n"))
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def first_value(data, keys):
    """Return the first non-empty value found among keys in a dict-like object."""
    if not isinstance(data, dict):
        return ""
    for k in keys:
        v = data.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return ""


OCR_FIELD_ORDER = [
    "verified_ocr", "verified_text", "ocr_text", "raw_text", "text", "content", "extracted_text",
]


def normalize_ocr_result(raw_result, fallback_text=""):
    """Accepts whatever the OCR backend returned (dict/str/object) and
    produces a normalized text string using OCR_FIELD_ORDER priority."""
    if raw_result is None:
        return fallback_text
    if isinstance(raw_result, str):
        return raw_result if raw_result.strip() else fallback_text
    if isinstance(raw_result, dict):
        val = first_value(raw_result, OCR_FIELD_ORDER)
        return val if val else fallback_text
    # object with attributes
    for k in OCR_FIELD_ORDER:
        v = getattr(raw_result, k, None)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return fallback_text


# --------------------------- OCR: Tesseract pipeline ------------------------
def _preprocess_variants(pil_img):
    """Yield several preprocessed versions of the image to try OCR on."""
    img = ImageOps.exif_transpose(pil_img).convert("RGB")

    # upscale small images for better OCR accuracy
    max_dim = max(img.size)
    if max_dim < 1600:
        scale = 1600 / max_dim
        img = img.resize((int(img.width * scale), int(img.height * scale)), Image.LANCZOS)

    gray = ImageOps.grayscale(img)

    # variant 1: plain grayscale
    yield gray

    # variant 2: autocontrast + sharpen
    v2 = ImageOps.autocontrast(gray)
    v2 = v2.filter(ImageFilter.SHARPEN)
    yield v2

    # variant 3: stronger contrast + threshold-ish
    v3 = ImageEnhance.Contrast(gray).enhance(2.0)
    v3 = v3.filter(ImageFilter.MedianFilter(size=3))
    yield v3


def tesseract_ocr(filepath, lang_hint="English"):
    """Run tesseract across several preprocessing variants and PSM modes,
    return the best (longest, highest-confidence) result."""
    if not TESSERACT_OK:
        return {"text": "", "confidence": 0.0}

    lang_map = {
        "English": "eng", "Tamil": "tam", "Hindi": "hin",
        "Telugu": "tel", "Kannada": "kan", "Malayalam": "mal",
    }
    tess_lang = lang_map.get(lang_hint, "eng")

    try:
        base_img = Image.open(filepath)
    except Exception as e:
        log.warning(f"Could not open image for OCR: {e}")
        return {"text": "", "confidence": 0.0}

    best_text = ""
    best_conf = 0.0
    psm_modes = [6, 4, 11, 3]  # block, columns, sparse, auto

    for variant in _preprocess_variants(base_img):
        for psm in psm_modes:
            cfg = f"--oem 3 --psm {psm}"
            try:
                # try requested language, fall back to eng if language pack missing
                try:
                    data = pytesseract.image_to_data(
                        variant, lang=tess_lang, config=cfg,
                        output_type=pytesseract.Output.DICT
                    )
                except pytesseract.TesseractError:
                    data = pytesseract.image_to_data(
                        variant, lang="eng", config=cfg,
                        output_type=pytesseract.Output.DICT
                    )

                words = []
                confs = []
                for i, w in enumerate(data.get("text", [])):
                    w = w.strip()
                    if not w:
                        continue
                    words.append(w)
                    try:
                        c = float(data["conf"][i])
                        if c >= 0:
                            confs.append(c)
                    except (ValueError, KeyError):
                        pass

                text = " ".join(words).strip()
                avg_conf = (sum(confs) / len(confs)) if confs else 0.0

                # score favours longer, higher-confidence extractions
                score = len(text) * (0.5 + avg_conf / 100.0)
                best_score = len(best_text) * (0.5 + best_conf / 100.0)
                if score > best_score:
                    best_text, best_conf = text, avg_conf
            except Exception as e:
                log.warning(f"tesseract variant/psm {psm} failed: {e}")
                continue

    return {"text": best_text.strip(), "confidence": round(best_conf, 1)}


def run_ocr(filepath, ext, lang_hint="English"):
    """Master OCR entrypoint: project OCR service first, tesseract fallback."""
    text = ""
    confidence = 0.0
    source = "none"

    if ext in IMAGE_EXTENSIONS:
        if PROJECT_OCR_FN:
            try:
                raw = PROJECT_OCR_FN(filepath)
                text = normalize_ocr_result(raw)
                if isinstance(raw, dict):
                    confidence = float(raw.get("confidence", raw.get("ocr_confidence", 0)) or 0)
                source = "project_service"
            except Exception as e:
                log.warning(f"Project OCR service failed, falling back to tesseract: {e}")
                text = ""

        if not text.strip():
            result = tesseract_ocr(filepath, lang_hint)
            text = result["text"]
            confidence = result["confidence"]
            source = "tesseract"

    elif ext == "pdf":
        text, confidence, source = extract_pdf_text(filepath, lang_hint)

    elif ext == "docx":
        text, confidence, source = extract_docx_text(filepath)

    text = clean_ocr_text(text)
    return text, confidence, source


def extract_pdf_text(filepath, lang_hint="English"):
    text_parts = []
    if PYPDF_OK:
        try:
            reader = PdfReader(filepath)
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    text_parts.append(t)
        except Exception as e:
            log.warning(f"pypdf extraction failed: {e}")

    text = "\n\n".join(text_parts).strip()
    if text:
        return text, 90.0, "pypdf"

    # image-based / scanned PDF with no extractable text: attempt OCR on
    # rendered pages if pypdfium2 is available
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(filepath)
        collected = []
        confs = []
        for i in range(len(pdf)):
            page = pdf[i]
            bitmap = page.render(scale=2.0)
            pil_img = bitmap.to_pil()
            tmp_path = os.path.join(PROCESSED_DIR, f"_pdf_page_{uuid.uuid4().hex}.png")
            pil_img.save(tmp_path)
            try:
                res = tesseract_ocr(tmp_path, lang_hint)
                if res["text"].strip():
                    collected.append(res["text"])
                    confs.append(res["confidence"])
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        text = "\n\n".join(collected).strip()
        avg_conf = round(sum(confs) / len(confs), 1) if confs else 0.0
        return text, avg_conf, "pdf_ocr"
    except Exception as e:
        log.warning(f"PDF OCR fallback unavailable/failed: {e}")
        return "", 0.0, "pdf_failed"


def extract_docx_text(filepath):
    if not DOCX_OK:
        return "", 0.0, "docx_unavailable"
    try:
        d = docx_lib.Document(filepath)
        paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paras).strip()
        return text, 95.0 if text else 0.0, "python-docx"
    except Exception as e:
        log.warning(f"docx extraction failed: {e}")
        return "", 0.0, "docx_failed"


# --------------------------- Local note-processing pipeline -----------------
STOPWORDS = set("""
a an the is are was were be been being of to in on at for with and or but if then
than so as by from into that this these those it its it's i you he she they we
his her their our your not no do does did done can could will would shall should
may might must have has had having about above after again against all am any
because before below between both down during each few further here how i'm
i've i'll just more most other over own same some such under until up very what
when where which who whom why itself myself yourself ourselves note notes noted
""".split())


def split_sentences(text):
    sentences = re.split(r"(?<=[.!?])\s+|\n+", text)
    return [s.strip() for s in sentences if s.strip() and len(s.strip()) > 3]


def extract_keywords(text, top_n=10):
    words = re.findall(r"[A-Za-z][A-Za-z\-']{2,}", text)
    freq = Counter()
    for w in words:
        lw = w.lower()
        if lw in STOPWORDS:
            continue
        freq[w if w.isupper() else lw] += 1
    # prefer capitalized / acronym terms as they read as "key terms"
    ranked = sorted(freq.items(), key=lambda kv: (-kv[1], -len(kv[0])))
    seen = set()
    keywords = []
    for word, _ in ranked:
        norm = word.lower()
        if norm in seen:
            continue
        seen.add(norm)
        keywords.append(word if word.isupper() else word.capitalize())
        if len(keywords) >= top_n:
            break
    return keywords


def summarize_text(text, max_sentences=4):
    sentences = split_sentences(text)
    if not sentences:
        return ""
    if len(sentences) <= max_sentences:
        return " ".join(sentences)

    keywords = set(k.lower() for k in extract_keywords(text, top_n=15))
    scored = []
    for idx, s in enumerate(sentences):
        words = re.findall(r"[A-Za-z][A-Za-z\-']{2,}", s.lower())
        score = sum(1 for w in words if w in keywords)
        # slight bonus for earlier sentences (topic sentences)
        score += max(0, 2 - idx * 0.1)
        scored.append((score, idx, s))

    top = sorted(scored, key=lambda t: -t[0])[:max_sentences]
    top_sorted = sorted(top, key=lambda t: t[1])  # restore original order
    return " ".join(s for _, _, s in top_sorted)


def build_revision_notes(text, max_points=6):
    sentences = split_sentences(text)
    if not sentences:
        return []
    keywords = set(k.lower() for k in extract_keywords(text, top_n=15))
    scored = []
    for idx, s in enumerate(sentences):
        words = re.findall(r"[A-Za-z][A-Za-z\-']{2,}", s.lower())
        score = sum(1 for w in words if w in keywords)
        scored.append((score, idx, s))
    top = sorted(scored, key=lambda t: -t[0])[:max_points]
    top_sorted = sorted(top, key=lambda t: t[1])
    points = []
    for i, (_, _, s) in enumerate(top_sorted, start=1):
        points.append({"title": f"Key Point {i}", "text": s})
    return points


def build_quiz(text, concepts, max_questions=5):
    sentences = split_sentences(text)
    quiz = []
    used = set()
    for concept in concepts:
        if len(quiz) >= max_questions:
            break
        c_lower = concept.lower()
        for s in sentences:
            if c_lower in s.lower() and s not in used:
                used.add(s)
                question = f"What does the note say about \"{concept}\"?"
                answer = s.strip()
                quiz.append({"question": question, "answer": answer, "concept": concept})
                break
    # if not enough concept-based questions, fill with fill-in-the-blank style
    if len(quiz) < max_questions:
        for s in sentences:
            if s in used or len(quiz) >= max_questions:
                continue
            words = re.findall(r"[A-Za-z][A-Za-z\-']{3,}", s)
            candidates = [w for w in words if w.lower() not in STOPWORDS]
            if not candidates:
                continue
            target = max(candidates, key=len)
            blanked = re.sub(r"\b" + re.escape(target) + r"\b", "_____", s, count=1)
            if blanked == s:
                continue
            used.add(s)
            quiz.append({
                "question": f"Fill in the blank: {blanked}",
                "answer": target,
                "concept": target,
            })
    return quiz[:max_questions]


def process_note_locally(text):
    text = text.strip()
    if not text:
        return {
            "summary": "",
            "concepts": [],
            "revision": [],
            "quiz": [],
        }
    concepts = extract_keywords(text, top_n=10)
    summary = summarize_text(text, max_sentences=4)
    revision = build_revision_notes(text, max_points=6)
    quiz = build_quiz(text, concepts, max_questions=5)
    return {
        "summary": summary,
        "concepts": concepts,
        "revision": revision,
        "quiz": quiz,
    }


def process_note(text):
    """Use project pipeline if available, else local fallback. Always returns
    a fully-populated dict so the frontend never sees missing fields."""
    if PROJECT_PIPELINE_FN:
        try:
            result = PROJECT_PIPELINE_FN(text)
            if isinstance(result, dict) and result.get("summary"):
                # backfill any missing fields with local processing
                local = process_note_locally(text)
                for k, v in local.items():
                    result.setdefault(k, v)
                return result
        except Exception as e:
            log.warning(f"Project note pipeline failed, using local fallback: {e}")
    return process_note_locally(text)


# --------------------------- Translation -------------------------------------
def translate_text(text, source, target):
    if not text.strip() or source == target:
        return text, True, "no-op"
    if TRANSLATOR_OK:
        try:
            lang_code = {
                "English": "en", "Tamil": "ta", "Hindi": "hi",
                "Telugu": "te", "Kannada": "kn", "Malayalam": "ml",
            }
            src = lang_code.get(source, "auto")
            tgt = lang_code.get(target, "en")
            translated = GoogleTranslator(source=src, target=tgt).translate(text)
            if translated and translated.strip():
                return translated, True, "deep_translator"
        except Exception as e:
            log.warning(f"Translation failed, returning original text: {e}")
    # graceful fallback: NEVER destroy original content
    return text, False, "unavailable"


# --------------------------- TTS ----------------------------------------------
def generate_tts(text, language):
    if not text.strip():
        return None, False, "empty_text"
    if GTTS_OK:
        try:
            lang_code = {
                "English": "en", "Tamil": "ta", "Hindi": "hi",
                "Telugu": "te", "Kannada": "kn", "Malayalam": "ml",
            }
            code = lang_code.get(language, "en")
            filename = f"{uuid.uuid4().hex}.mp3"
            path = os.path.join(AUDIO_DIR, filename)
            tts = gTTS(text=text[:4500], lang=code)
            tts.save(path)
            return filename, True, "gtts"
        except Exception as e:
            log.warning(f"gTTS failed: {e}")
    return None, False, "server_tts_unavailable"


# ---------------------------------------------------------------------------
# ROUTES
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api")
def api_root():
    return jsonify({
        "success": True,
        "name": "NoteWise AI API",
        "version": "1.0",
        "endpoints": [
            "/api/health", "/api/languages", "/api/process-upload",
            "/api/process-text", "/api/translate", "/api/tts",
            "/api/session", "/api/audio/<filename>",
        ],
    })


@app.route("/api/health")
def health():
    return jsonify({
        "success": True,
        "status": "online",
        "ocr": bool(OCR_AVAILABLE and (TESSERACT_OK or PROJECT_OCR_FN)),
        "note_pipeline": bool(NOTE_PIPELINE_AVAILABLE),
        "project_ocr_service": bool(PROJECT_OCR_FN),
        "project_note_pipeline": bool(PROJECT_PIPELINE_FN),
        "tesseract": TESSERACT_OK,
        "pdf_support": PYPDF_OK,
        "docx_support": DOCX_OK,
        "server_tts": GTTS_OK,
        "translation_service": TRANSLATOR_OK,
    })


@app.route("/api/languages")
def languages():
    return jsonify({"success": True, "languages": LANGUAGES})


@app.route("/api/process-upload", methods=["POST"])
def process_upload():
    try:
        if "file" not in request.files:
            return jsonify({"success": False, "error": "No file was uploaded."}), 400

        file = request.files["file"]
        if file.filename == "":
            return jsonify({"success": False, "error": "No file was selected."}), 400

        if not allowed_file(file.filename):
            return jsonify({
                "success": False,
                "error": "Unsupported file type. Please upload JPG, PNG, WEBP, BMP, TIFF, PDF or DOCX.",
            }), 400

        language = request.form.get("language", "English")
        input_language = request.form.get("input_language", request.form.get("language", "English"))

        ext = file.filename.rsplit(".", 1)[1].lower()
        saved_name = unique_filename(file.filename)
        saved_path = os.path.join(UPLOAD_DIR, saved_name)
        file.save(saved_path)

        text, confidence, source = run_ocr(saved_path, ext, input_language)

        if not text.strip():
            return jsonify({
                "success": False,
                "error": (
                    "No readable text could be extracted from this file. "
                    "Try a clearer photo, better lighting, or a higher-resolution scan."
                ),
                "ocr_source": source,
            }), 200

        processed = process_note(text)

        response = {
            "success": True,
            "raw_text": text,
            "ocr_text": text,
            "verified_ocr": text,
            "verified_text": text,
            "note_content": text,
            "summary": processed.get("summary", ""),
            "concepts": processed.get("concepts", []),
            "revision": processed.get("revision", []),
            "quiz": processed.get("quiz", []),
            "confidence": confidence,
            "ocr_confidence": confidence,
            "ocr_source": source,
            "input_language": input_language,
            "output_language": language,
            "filename": file.filename,
        }
        return jsonify(response)

    except Exception as e:
        log.error(f"process_upload failed: {e}\n{traceback.format_exc()}")
        return jsonify({"success": False, "error": f"Server error while processing file: {e}"}), 500


@app.route("/api/process-text", methods=["POST"])
def process_text():
    try:
        data = request.get_json(force=True, silent=True) or {}
        text = (data.get("text") or "").strip()
        if not text:
            return jsonify({"success": False, "error": "No text provided."}), 400

        text = clean_ocr_text(text)
        processed = process_note(text)
        return jsonify({
            "success": True,
            "raw_text": text,
            "ocr_text": text,
            "verified_ocr": text,
            "verified_text": text,
            "note_content": text,
            "summary": processed.get("summary", ""),
            "concepts": processed.get("concepts", []),
            "revision": processed.get("revision", []),
            "quiz": processed.get("quiz", []),
            "confidence": 100.0,
            "ocr_confidence": 100.0,
        })
    except Exception as e:
        log.error(f"process_text failed: {e}\n{traceback.format_exc()}")
        return jsonify({"success": False, "error": f"Server error: {e}"}), 500


@app.route("/api/translate", methods=["POST"])
def translate():
    try:
        data = request.get_json(force=True, silent=True) or {}
        text = (data.get("text") or "").strip()
        source = data.get("source", "English")
        target = data.get("target", "English")

        if not text:
            return jsonify({"success": False, "error": "No text provided."}), 400

        translated, real_translation, engine = translate_text(text, source, target)
        return jsonify({
            "success": True,
            "text": translated,
            "engine": engine,
            "note": None if real_translation else (
                "Live translation service is unavailable right now, so the "
                "original text is shown unchanged."
            ),
        })
    except Exception as e:
        log.error(f"translate failed: {e}\n{traceback.format_exc()}")
        return jsonify({"success": False, "error": f"Server error: {e}"}), 500


@app.route("/api/tts", methods=["POST"])
def tts():
    try:
        data = request.get_json(force=True, silent=True) or {}
        text = (data.get("text") or "").strip()
        language = data.get("language", "English")

        if not text:
            return jsonify({"success": False, "error": "No text provided."}), 400

        filename, ok, engine = generate_tts(text, language)
        if ok and filename:
            return jsonify({
                "success": True,
                "audio_url": url_for("get_audio", filename=filename),
                "engine": engine,
            })
        return jsonify({
            "success": False,
            "error": "Server-side speech generation is unavailable. Using browser speech instead.",
            "fallback": "browser_speech",
        })
    except Exception as e:
        log.error(f"tts failed: {e}\n{traceback.format_exc()}")
        return jsonify({"success": False, "error": f"Server error: {e}", "fallback": "browser_speech"}), 200


@app.route("/api/audio/<path:filename>")
def get_audio(filename):
    safe = secure_filename(filename)
    return send_from_directory(AUDIO_DIR, safe, mimetype="audio/mpeg")


@app.route("/api/session", methods=["POST"])
def session():
    try:
        return jsonify({"success": True, "session_id": uuid.uuid4().hex})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.errorhandler(413)
def too_large(e):
    return jsonify({"success": False, "error": "File is too large. Maximum upload size is 25MB."}), 413


@app.errorhandler(404)
def not_found(e):
    return jsonify({"success": False, "error": "Not found."}), 404


@app.errorhandler(500)
def server_error(e):
    return jsonify({"success": False, "error": "Internal server error."}), 500


# ---------------------------------------------------------------------------
# STARTUP HELPERS
# ---------------------------------------------------------------------------
def kill_port(port):
    """Best-effort: free the port before starting (Linux only, ignores errors)."""
    try:
        os.system(f"fuser -k {port}/tcp >/dev/null 2>&1")
    except Exception:
        pass


def print_status_banner():
    print("=" * 60)
    print("NOTEWISE AI - starting")
    print("=" * 60)
    print(f"OCR (tesseract):        {'AVAILABLE' if TESSERACT_OK else 'UNAVAILABLE'}")
    print(f"Project OCR service:    {'FOUND' if PROJECT_OCR_FN else 'not found (using tesseract)'}")
    print(f"Project note pipeline:  {'FOUND' if PROJECT_PIPELINE_FN else 'not found (using local fallback)'}")
    print(f"PDF support (pypdf):    {'AVAILABLE' if PYPDF_OK else 'UNAVAILABLE'}")
    print(f"DOCX support:           {'AVAILABLE' if DOCX_OK else 'UNAVAILABLE'}")
    print(f"Server TTS (gTTS):      {'AVAILABLE' if GTTS_OK else 'UNAVAILABLE (browser speech fallback active)'}")
    print(f"Translation service:    {'AVAILABLE' if TRANSLATOR_OK else 'UNAVAILABLE (fallback echoes original text)'}")
    print("=" * 60)


if __name__ == "__main__":
    kill_port(7860)
    print_status_banner()
    app.run(host="0.0.0.0", port=7860, debug=False, use_reloader=False, threaded=True)
